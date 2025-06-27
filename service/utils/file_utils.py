import PyPDF2
import docx
import os
import logging
from pathlib import Path
from typing import List, Optional, Dict, Any

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class PDFProcessor:
    """
    Utility class for processing PDF documents
    """
    
    def extract_text(self, pdf_path: str) -> str:
        """
        Extract text content from a PDF file
        
        Args:
            pdf_path: Path to the PDF file
            
        Returns:
            Extracted text content
        """
        try:
            text = ""
            with open(pdf_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    text += page.extract_text() + "\n"
            
            logger.info(f"Successfully extracted text from {pdf_path}")
            return text.strip()
            
        except Exception as e:
            logger.error(f"Error extracting text from PDF {pdf_path}: {str(e)}")
            return ""
    
    def extract_metadata(self, pdf_path: str) -> Dict[str, Any]:
        """
        Extract metadata from a PDF file
        
        Args:
            pdf_path: Path to the PDF file
            
        Returns:
            Dictionary containing metadata
        """
        try:
            with open(pdf_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                metadata = pdf_reader.metadata
                
                return {
                    "title": metadata.get("/Title", ""),
                    "author": metadata.get("/Author", ""),
                    "subject": metadata.get("/Subject", ""),
                    "creator": metadata.get("/Creator", ""),
                    "producer": metadata.get("/Producer", ""),
                    "creation_date": metadata.get("/CreationDate", ""),
                    "modification_date": metadata.get("/ModDate", ""),
                    "pages": len(pdf_reader.pages)
                }
        except Exception as e:
            logger.error(f"Error extracting metadata from PDF {pdf_path}: {str(e)}")
            return {}

class DocumentProcessor:
    """
    Utility class for processing various document formats
    """
    
    def __init__(self):
        self.pdf_processor = PDFProcessor()
    
    def process_document(self, file_path: str) -> Dict[str, Any]:
        """
        Process a document and extract its content
        
        Args:
            file_path: Path to the document file
            
        Returns:
            Dictionary containing document content and metadata
        """
        file_path = Path(file_path)
        
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        file_extension = file_path.suffix.lower()
        
        if file_extension == '.pdf':
            return self._process_pdf(str(file_path))
        elif file_extension in ['.docx', '.doc']:
            return self._process_word(str(file_path))
        elif file_extension == '.txt':
            return self._process_text(str(file_path))
        else:
            raise ValueError(f"Unsupported file format: {file_extension}")
    
    def _process_pdf(self, file_path: str) -> Dict[str, Any]:
        """Process PDF document"""
        text = self.pdf_processor.extract_text(file_path)
        metadata = self.pdf_processor.extract_metadata(file_path)
        
        return {
            "content": text,
            "metadata": metadata,
            "file_type": "pdf",
            "file_path": file_path
        }
    
    def _process_word(self, file_path: str) -> Dict[str, Any]:
        """Process Word document"""
        try:
            doc = docx.Document(file_path)
            text = ""
            
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Extract basic metadata
            core_props = doc.core_properties
            metadata = {
                "title": core_props.title or "",
                "author": core_props.author or "",
                "subject": core_props.subject or "",
                "created": str(core_props.created) if core_props.created else "",
                "modified": str(core_props.modified) if core_props.modified else "",
                "paragraphs": len(doc.paragraphs)
            }
            
            return {
                "content": text.strip(),
                "metadata": metadata,
                "file_type": "docx",
                "file_path": file_path
            }
            
        except Exception as e:
            logger.error(f"Error processing Word document {file_path}: {str(e)}")
            return {
                "content": "",
                "metadata": {},
                "file_type": "docx",
                "file_path": file_path,
                "error": str(e)
            }
    
    def _process_text(self, file_path: str) -> Dict[str, Any]:
        """Process plain text document"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                text = file.read()
            
            # Get file stats
            stat = os.stat(file_path)
            metadata = {
                "size": stat.st_size,
                "created": stat.st_ctime,
                "modified": stat.st_mtime,
                "lines": len(text.splitlines())
            }
            
            return {
                "content": text,
                "metadata": metadata,
                "file_type": "txt",
                "file_path": file_path
            }
            
        except Exception as e:
            logger.error(f"Error processing text file {file_path}: {str(e)}")
            return {
                "content": "",
                "metadata": {},
                "file_type": "txt",
                "file_path": file_path,
                "error": str(e)
            }
    
    def batch_process_documents(self, directory_path: str, 
                              file_extensions: Optional[List[str]] = None) -> List[Dict[str, Any]]:
        """
        Process multiple documents in a directory
        
        Args:
            directory_path: Path to the directory containing documents
            file_extensions: List of file extensions to process (default: ['.pdf', '.docx', '.txt'])
            
        Returns:
            List of processed document dictionaries
        """
        if file_extensions is None:
            file_extensions = ['.pdf', '.docx', '.txt']
        
        directory = Path(directory_path)
        if not directory.exists() or not directory.is_dir():
            raise ValueError(f"Invalid directory path: {directory_path}")
        
        processed_docs = []
        
        for file_path in directory.rglob('*'):
            if file_path.is_file() and file_path.suffix.lower() in file_extensions:
                try:
                    doc_data = self.process_document(str(file_path))
                    processed_docs.append(doc_data)
                    logger.info(f"Processed document: {file_path}")
                except Exception as e:
                    logger.error(f"Error processing document {file_path}: {str(e)}")
        
        return processed_docs

class FileValidator:
    """
    Utility class for validating file integrity and format
    """
    
    @staticmethod
    def validate_file_size(file_path: str, max_size_mb: float = 50.0) -> bool:
        """
        Validate file size is within acceptable limits
        
        Args:
            file_path: Path to the file
            max_size_mb: Maximum file size in MB
            
        Returns:
            True if file size is acceptable, False otherwise
        """
        try:
            file_size = os.path.getsize(file_path)
            max_size_bytes = max_size_mb * 1024 * 1024
            return file_size <= max_size_bytes
        except Exception:
            return False
    
    @staticmethod
    def validate_file_format(file_path: str, allowed_formats: List[str]) -> bool:
        """
        Validate file format is in allowed list
        
        Args:
            file_path: Path to the file
            allowed_formats: List of allowed file extensions
            
        Returns:
            True if file format is allowed, False otherwise
        """
        try:
            file_extension = Path(file_path).suffix.lower()
            return file_extension in [fmt.lower() for fmt in allowed_formats]
        except Exception:
            return False
    
    @staticmethod
    def is_file_readable(file_path: str) -> bool:
        """
        Check if file is readable
        
        Args:
            file_path: Path to the file
            
        Returns:
            True if file is readable, False otherwise
        """
        try:
            with open(file_path, 'rb') as file:
                file.read(1)  # Try to read one byte
            return True
        except Exception:
            return False
